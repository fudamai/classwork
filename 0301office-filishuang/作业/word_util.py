#! python
# -*- coding:utf-8 -*-
# Author:fudamai
# word_util.py


from docx import Document


class Word_util:
    def __init__(self, text,in_doc, w_type='save'):  # 
        "定制实例属性"
        self.in_text = text
        self.w_type = w_type
        self.doc = in_doc

    def add_text(self):
        """根据输入的参数判断输入文本的格式，并调用相应的函数"""
        if self.w_type == 'title':
            self.add_title()
        elif self.w_type == 'subtitle':
            self.add_subtitle()
        elif self.w_type == 'para':
            self.add_para()
        elif self.w_type == 'save':
            self.save()

    def save(self):
        "将输入内容进行保存，自动添加.doc后缀"
        self.doc.save(f'{self.in_text}.doc')
        

    def add_title(self):
        "添加头部标题"
        self.doc.add_heading(self.in_text)
        # self.doc.save('quit.doc')
        return self.doc
        

    def add_subtitle(self):
        "添加副标题"
        self.doc.add_paragraph(self.in_text,'Subtitle')
        # self.doc.save('quit.doc')
        return self.doc
        
    
    def add_para(self):
        "添加正文"
        self.doc.add_paragraph(self.in_text)
        # self.doc.save('quit.doc')
        return self.doc
        


def main():
    print('欢迎使用Word'.center(30,'-'))
    menu ={
        'title':'选择输入标题',
        'subtitle':'选择输入副标题',
        'para':'选择输入正文',
        'save':'输入文件名并保存',
        'quit':'关闭程序'
    }
    for k,v in menu.items():
        print(f'{k}:{v}')


    w_type = input('请根据需要输入命令：')
    in_doc = Document()  # 在类中指定实例会导致每次循环生成新的实例，调用save时保存的是空文件
    while w_type: 
        if w_type == 'quit':
            print('即将关闭程序')
            break
        
        text= input('输入文本：')
        Word_util(text=text, in_doc=in_doc, w_type=w_type).add_text()  # 使用关键字参数
        w_type = input('请根据需要输入命令：')

if __name__ == "__main__":
    main()       